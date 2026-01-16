"""Service for processing RAG documents (files and URLs)"""
import os
import aiofiles
import aiohttp
from typing import Dict, Any, Optional
from pathlib import Path
import mimetypes
from urllib.parse import urlparse
import logging
import string

logger = logging.getLogger(__name__)

# Create uploads directory if it doesn't exist
UPLOADS_DIR = Path("uploads/rag_documents")
UPLOADS_DIR.mkdir(parents=True, exist_ok=True)


async def save_uploaded_file(file_content: bytes, filename: str) -> Dict[str, Any]:
    """Save uploaded file to disk and return file info"""
    try:
        # Generate unique filename
        file_ext = Path(filename).suffix
        import uuid
        unique_filename = f"{uuid.uuid4()}{file_ext}"
        file_path = UPLOADS_DIR / unique_filename
        
        # Save file
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(file_content)
        
        # Get file info
        file_size = len(file_content)
        mime_type, _ = mimetypes.guess_type(filename)
        
        # Improve MIME type detection for common Office formats
        file_ext = Path(filename).suffix.lower()
        if not mime_type:
            # Manual MIME type mapping for common formats
            mime_type_map = {
                '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
                '.doc': 'application/msword',
                '.xlsx': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                '.xls': 'application/vnd.ms-excel',
                '.pptx': 'application/vnd.openxmlformats-officedocument.presentationml.presentation',
                '.ppt': 'application/vnd.ms-powerpoint',
            }
            mime_type = mime_type_map.get(file_ext, "application/octet-stream")
        
        return {
            "file_path": str(file_path),
            "file_name": filename,
            "file_type": mime_type,
            "file_size": file_size,
            "unique_filename": unique_filename
        }
    except Exception as e:
        logger.error(f"Error saving file: {e}")
        raise


async def fetch_url_content(url: str) -> Dict[str, Any]:
    """Fetch content from URL and extract text"""
    try:
        async with aiohttp.ClientSession() as session:
            async with session.get(url, timeout=aiohttp.ClientTimeout(total=30)) as response:
                if response.status != 200:
                    raise Exception(f"Failed to fetch URL: HTTP {response.status}")
                
                content_type = response.headers.get('Content-Type', '')
                
                # Check if it's HTML
                if 'text/html' in content_type:
                    html_content = await response.text()
                    # Basic HTML text extraction (can be enhanced with BeautifulSoup)
                    import re
                    # Remove script and style elements
                    html_content = re.sub(r'<script[^>]*>.*?</script>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                    html_content = re.sub(r'<style[^>]*>.*?</style>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                    # Extract text from HTML tags
                    text_content = re.sub(r'<[^>]+>', ' ', html_content)
                    text_content = re.sub(r'\s+', ' ', text_content).strip()
                    
                    return {
                        "content_text": text_content[:50000],  # Limit to 50k chars
                        "content_metadata": {
                            "url": url,
                            "content_type": content_type,
                            "content_length": len(text_content)
                        }
                    }
                else:
                    # Try to get as text
                    text_content = await response.text()
                    return {
                        "content_text": text_content[:50000],
                        "content_metadata": {
                            "url": url,
                            "content_type": content_type,
                            "content_length": len(text_content)
                        }
                    }
    except Exception as e:
        logger.error(f"Error fetching URL: {e}")
        raise


async def extract_text_from_file(file_path: str, file_type: str) -> Dict[str, Any]:
    """Extract text content from various file types"""
    try:
        content_text = ""
        content_metadata = {}
        
        # Text files
        if file_type.startswith('text/'):
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content_text = await f.read()
        
        # PDF files
        elif file_type == 'application/pdf':
            try:
                import PyPDF2
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    pages = []
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        pages.append(text)
                        content_text += text + "\n"
                    content_metadata = {"page_count": len(pdf_reader.pages)}
            except ImportError:
                content_text = "[PDF processing requires PyPDF2 library]"
                logger.warning("PyPDF2 not installed, cannot extract PDF text")
            except Exception as e:
                logger.error(f"Error extracting PDF text: {e}")
                raise
        
        # Word documents (check by file extension too, in case MIME type detection fails)
        elif (file_type in ['application/msword', 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'] 
              or str(file_path).lower().endswith(('.doc', '.docx'))):
            logger.info(f"Processing Word document: {file_path}, detected type: {file_type}")
            
            # CRITICAL: First verify file is not corrupted/binary by checking first bytes
            try:
                async with aiofiles.open(file_path, 'rb') as f:
                    first_bytes = await f.read(4)
                    # Check for ZIP signature (DOCX is a ZIP archive)
                    if len(first_bytes) >= 4 and first_bytes[:2] == b'PK':
                        if len(first_bytes) >= 4 and (first_bytes[2] == 3 and first_bytes[3] == 4) or (first_bytes[2] == 5 and first_bytes[3] == 6):
                            # File is a valid ZIP/DOCX file, safe to process with python-docx
                            logger.info(f"DOCX file signature verified. Proceeding with text extraction...")
                        else:
                            logger.warning(f"File has PK signature but invalid format. May not be a valid DOCX file.")
                    else:
                        logger.warning(f"File does not appear to be a valid DOCX/ZIP file. Signature: {first_bytes[:4]}")
            except Exception as sig_check_error:
                logger.error(f"Error checking DOCX file signature: {sig_check_error}")
                # Continue anyway - python-docx will handle it
            
            try:
                from docx import Document
                logger.info(f"Attempting to extract text from DOCX file: {file_path}")
                
                # Try to open and read the DOCX file
                try:
                    doc = Document(file_path)
                except Exception as doc_error:
                    logger.error(f"Failed to open DOCX file {file_path}: {doc_error}", exc_info=True)
                    # CRITICAL: Don't let binary content leak through
                    content_text = f"[Error opening DOCX file: {str(doc_error)}. File may be corrupted.]"
                    content_metadata = {"error": str(doc_error), "file_type": file_type}
                    raise
                
                paragraphs = []
                content_text = ""  # Initialize to empty string
                for para in doc.paragraphs:
                    if para.text and para.text.strip():
                        # CRITICAL: Validate each paragraph is text, not binary
                        para_text = para.text.strip()
                        if para_text and not para_text.startswith('PK'):  # Quick check
                            paragraphs.append(para_text)
                            content_text += para_text + "\n"
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_cells = []
                        for cell in row.cells:
                            if cell.text and cell.text.strip():
                                cell_text = cell.text.strip()
                                # Quick validation
                                if not cell_text.startswith('PK'):
                                    row_cells.append(cell_text)
                        if row_cells:
                            row_text = " | ".join(row_cells)
                            if row_text:
                                content_text += row_text + "\n"
                
                logger.info(f"Successfully extracted {len(paragraphs)} paragraphs from DOCX file, total text length: {len(content_text)}")
                
                # CRITICAL: Ensure we have actual text content, not binary
                if not content_text or len(content_text.strip()) == 0:
                    content_text = "[No text content found in Word document]"
                    logger.warning(f"No text content extracted from DOCX file: {file_path}")
                else:
                    # CRITICAL: Validate that content_text is actually text, not binary
                    # Check for ZIP signature (should NEVER be in extracted text)
                    if isinstance(content_text, str):
                        # Check first 4 bytes for ZIP signature
                        first_chars = content_text[:4] if len(content_text) >= 4 else content_text
                        if first_chars.startswith('PK'):
                            # Check if it's actually a ZIP signature (PK\x03\x04)
                            if len(first_chars) >= 4:
                                char2 = ord(first_chars[2]) if len(first_chars) > 2 else 0
                                char3 = ord(first_chars[3]) if len(first_chars) > 3 else 0
                                if (char2 == 3 and char3 == 4) or (char2 == 5 and char3 == 6):
                                    logger.error(f"CRITICAL: Binary ZIP content detected in DOCX extraction result! File: {file_path}")
                                    content_text = "[Error: Binary content detected after text extraction. DOCX file may be corrupted or extraction failed.]"
                        # Also check for null bytes (should never be in text)
                        if '\x00' in content_text:
                            null_count = content_text.count('\x00')
                            logger.error(f"CRITICAL: Found {null_count} null bytes in DOCX extraction result! File: {file_path}")
                            content_text = content_text.replace('\x00', '')
                            # If still has ZIP signature after removing nulls, reject entirely
                            if content_text.startswith('PK') and len(content_text) >= 4:
                                char2 = ord(content_text[2]) if len(content_text) > 2 else 0
                                char3 = ord(content_text[3]) if len(content_text) > 3 else 0
                                if (char2 == 3 and char3 == 4) or (char2 == 5 and char3 == 6):
                                    logger.error(f"CRITICAL: ZIP signature still present after null byte removal! Rejecting entirely.")
                                    content_text = "[Error: Binary content could not be cleaned. DOCX file extraction failed.]"
                    elif isinstance(content_text, bytes):
                        # This should never happen, but if it does, decode it
                        logger.error(f"CRITICAL: content_text is bytes instead of string after DOCX extraction!")
                        try:
                            content_text = content_text.decode('utf-8', errors='ignore')
                            # Check again after decoding
                            if content_text.startswith('PK') and len(content_text) >= 4:
                                char2 = ord(content_text[2]) if len(content_text) > 2 else 0
                                char3 = ord(content_text[3]) if len(content_text) > 3 else 0
                                if (char2 == 3 and char3 == 4) or (char2 == 5 and char3 == 6):
                                    content_text = "[Error: Binary content detected after DOCX extraction.]"
                        except:
                            content_text = "[Error: Unable to decode binary content from DOCX extraction.]"
                
                content_metadata = {"paragraph_count": len(paragraphs)}
                    
            except ImportError:
                content_text = "[Word document processing requires python-docx library]"
                logger.warning("python-docx not installed, cannot extract Word text")
                content_metadata = {"error": "python-docx library not installed"}
            except Exception as e:
                logger.error(f"Error extracting Word text from {file_path}: {e}", exc_info=True)
                # CRITICAL: Don't raise - return error message instead of binary content
                # NEVER return binary file content
                content_text = f"[Error extracting text from Word document: {str(e)}]"
                content_metadata = {"error": str(e)}
        
        # Images (OCR)
        elif file_type.startswith('image/'):
            try:
                from PIL import Image
                import pytesseract
                
                image = Image.open(file_path)
                content_text = pytesseract.image_to_string(image)
                content_metadata = {
                    "image_width": image.width,
                    "image_height": image.height,
                    "image_format": image.format
                }
            except ImportError:
                content_text = "[Image OCR requires PIL and pytesseract libraries]"
                logger.warning("PIL/pytesseract not installed, cannot extract image text")
            except Exception as e:
                logger.error(f"Error extracting image text: {e}")
                # Don't raise, just log - images without text are OK
                content_text = f"[Image file: {Path(file_path).name}]"
        
        # Excel files
        elif file_type in ['application/vnd.ms-excel', 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet']:
            try:
                import pandas as pd
                # Try to read as Excel
                df = pd.read_excel(file_path, sheet_name=None)
                sheets_text = []
                for sheet_name, sheet_df in df.items():
                    sheet_text = f"Sheet: {sheet_name}\n{sheet_df.to_string()}\n"
                    sheets_text.append(sheet_text)
                    content_text += sheet_text
                content_metadata = {"sheet_count": len(df), "sheet_names": list(df.keys())}
            except ImportError:
                content_text = "[Excel processing requires pandas and openpyxl libraries]"
                logger.warning("pandas/openpyxl not installed, cannot extract Excel text")
            except Exception as e:
                logger.error(f"Error extracting Excel text: {e}")
                raise
        
        # Default: try to read as text (only if it's actually a text file)
        else:
            # Check file extension to avoid trying to read binary files as text
            file_ext = Path(file_path).suffix.lower()
            # Known binary file extensions - don't try to read as text
            binary_extensions = {'.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx', 
                                '.zip', '.rar', '.7z', '.exe', '.dll', '.so', '.bin', '.img',
                                '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff', '.webp',
                                '.mp3', '.mp4', '.avi', '.mov', '.wmv', '.flv', '.mkv'}
            
            if file_ext in binary_extensions:
                content_text = f"[Binary file type not supported for direct text extraction: {Path(file_path).name}. Please use the appropriate file type handler.]"
                content_metadata = {"file_extension": file_ext, "message": "Binary file type detected"}
            else:
                # Try to read as text, but check for binary content
                try:
                    async with aiofiles.open(file_path, 'rb') as f:
                        raw_content = await f.read()
                    # Check if content contains null bytes (indicates binary)
                    if b'\x00' in raw_content[:1024]:  # Check first 1KB
                        content_text = f"[Binary file detected: {Path(file_path).name}]"
                        content_metadata = {"file_extension": file_ext, "message": "Binary content detected"}
                    else:
                        # Safe to decode as text
                        try:
                            content_text = raw_content.decode('utf-8', errors='ignore')
                        except UnicodeDecodeError:
                            content_text = f"[Unable to decode file as UTF-8: {Path(file_path).name}]"
                            content_metadata = {"file_extension": file_ext, "message": "Decoding error"}
                except Exception as e:
                    logger.error(f"Error reading file {file_path}: {e}")
                    content_text = f"[Error reading file: {str(e)}]"
                    content_metadata = {"error": str(e)}
        
        # Sanitize content_text to ensure it's valid UTF-8 and contains no null bytes
        if content_text:
            # If content_text is bytes, decode it
            if isinstance(content_text, bytes):
                try:
                    content_text = content_text.decode('utf-8', errors='ignore')
                except:
                    content_text = "[Unable to decode binary content as text]"
            
            # Check if content looks like binary (ZIP signature indicates DOCX/Office files)
            if isinstance(content_text, str):
                # CRITICAL: Check for ZIP file signature (DOCX files are ZIP archives) - check first 4 bytes
                # PK\x03\x04 is the local file header signature (ZIP/DOCX)
                # This should NEVER appear in extracted text - if it does, extraction failed
                if len(content_text) >= 4 and content_text.startswith('PK'):
                    char0 = ord(content_text[0])  # Should be 'P' = 80
                    char1 = ord(content_text[1])  # Should be 'K' = 75
                    char2 = ord(content_text[2]) if len(content_text) > 2 else 0
                    char3 = ord(content_text[3]) if len(content_text) > 3 else 0
                    
                    # ZIP signature: PK\x03\x04 or PK\x05\x06
                    if (char2 == 3 and char3 == 4) or (char2 == 5 and char3 == 6):
                        logger.error(f"CRITICAL: Binary ZIP/DOCX signature detected in content_text! First 50 chars: {repr(content_text[:50])}. Text extraction FAILED.")
                        content_text = "[Binary file content detected. Text extraction failed. DOCX file may be corrupted or python-docx extraction failed.]"
                    # Also check if content contains ZIP signature anywhere in first 200 chars (broader check)
                    elif '\x03\x04' in content_text[:200] or '\x05\x06' in content_text[:200]:
                        # Check if it's part of a ZIP signature pattern
                        pk_index = content_text.find('PK')
                        if pk_index >= 0 and pk_index < 200:
                            if pk_index + 2 < len(content_text):
                                next_char2 = ord(content_text[pk_index + 2])
                                next_char3 = ord(content_text[pk_index + 3]) if pk_index + 3 < len(content_text) else 0
                                if (next_char2 == 3 and next_char3 == 4) or (next_char2 == 5 and next_char3 == 6):
                                    logger.error(f"CRITICAL: ZIP signature pattern found in content_text at position {pk_index}! Text extraction FAILED.")
                                    content_text = "[Binary file content detected. Text extraction failed. DOCX file extraction failed.]"
                
                # Check for excessive null bytes (strong indicator of binary content)
                null_count = content_text.count('\x00')
                if null_count > 10:
                    logger.error(f"Excessive null bytes ({null_count}) detected in content_text! First 100 chars: {repr(content_text[:100])}")
                    content_text = "[Binary content detected. Text extraction may have failed.]"
                # Check if content has too many non-printable characters (indicates binary)
                elif len(content_text) > 0:
                    sample = content_text[:200] if len(content_text) > 200 else content_text
                    non_printable = sum(1 for c in sample if ord(c) < 32 and c not in '\n\r\t')
                    if non_printable > len(sample) * 0.3:  # More than 30% non-printable
                        logger.error(f"Too many non-printable characters ({non_printable}/{len(sample)}) in content_text! Likely binary content.")
                        content_text = "[Binary content detected. Text extraction may have failed.]"
            
            # CRITICAL: Remove null bytes and other invalid UTF-8 sequences BEFORE processing
            if isinstance(content_text, str):
                # Remove ALL null bytes first (PostgreSQL will reject them)
                if '\x00' in content_text:
                    null_count_before = content_text.count('\x00')
                    logger.warning(f"Removing {null_count_before} null bytes from content_text...")
                    content_text = content_text.replace('\x00', '')
                
                # CRITICAL: Check for ZIP signature BEFORE other processing
                if len(content_text) >= 4 and content_text.startswith('PK'):
                    char2 = ord(content_text[2]) if len(content_text) > 2 else 0
                    char3 = ord(content_text[3]) if len(content_text) > 3 else 0
                    if (char2 == 3 and char3 == 4) or (char2 == 5 and char3 == 6):
                        logger.error("CRITICAL: ZIP signature detected BEFORE cleaning! Rejecting entirely.")
                        content_text = "[Binary file content detected. Text extraction failed. Document contains binary ZIP data.]"
                        # Skip further processing for this error case
                    else:
                        # Remove other control characters except newline, carriage return, and tab
                        # Keep printable characters, newlines, carriage returns, and tabs
                        content_text = ''.join(c if c in string.printable or c in '\n\r\t' else ' ' for c in content_text)
                        # Limit length to 500k chars (increased from 100k to handle large documents like 93-page PDFs)
                        # This is still a reasonable limit for processing, and chunks will be created from this
                        content_text = content_text[:500000]
                        
                        # Final check: ensure no ZIP signature remains after cleaning
                        if len(content_text) >= 4 and content_text.startswith('PK'):
                            char2 = ord(content_text[2]) if len(content_text) > 2 else 0
                            char3 = ord(content_text[3]) if len(content_text) > 3 else 0
                            if (char2 == 3 and char3 == 4) or (char2 == 5 and char3 == 6):
                                logger.error("CRITICAL: ZIP signature still present after cleaning! Replacing with error message.")
                                content_text = "[Binary content could not be cleaned. Text extraction failed.]"
                else:
                    # No ZIP signature, safe to clean
                    # Remove other control characters except newline, carriage return, and tab
                    content_text = ''.join(c if c in string.printable or c in '\n\r\t' else ' ' for c in content_text)
                    # Limit length
                    content_text = content_text[:500000]
        else:
            content_text = ""
        
        # FINAL VALIDATION: Ensure content_text is safe for database storage
        # This is the LAST chance to catch binary content before it goes to the database
        if content_text and isinstance(content_text, str):
            # CRITICAL: Check for null bytes (PostgreSQL will reject these)
            if '\x00' in content_text:
                null_count = content_text.count('\x00')
                logger.error(f"CRITICAL: Found {null_count} null bytes in content_text at final validation! Removing ALL...")
                content_text = content_text.replace('\x00', '')
            
            # CRITICAL: Check for ZIP signature one final time
            if len(content_text) >= 4 and content_text.startswith('PK'):
                char2 = ord(content_text[2]) if len(content_text) > 2 else 0
                char3 = ord(content_text[3]) if len(content_text) > 3 else 0
                if (char2 == 3 and char3 == 4) or (char2 == 5 and char3 == 6):
                    logger.error(f"CRITICAL: ZIP signature detected in content_text at FINAL validation! File: {file_path}. Rejecting entirely.")
                    content_text = "[Binary file content detected. Text extraction failed. Document may be corrupted or extraction library failed.]"
                    content_metadata = {**content_metadata, "extraction_error": "Binary content detected at final validation"}
                # Also check for ZIP signature pattern elsewhere in first 200 chars
                elif len(content_text) > 200:
                    pk_pos = content_text.find('PK', 0, 200)
                    if pk_pos >= 0 and pk_pos + 3 < len(content_text):
                        next_char2 = ord(content_text[pk_pos + 2])
                        next_char3 = ord(content_text[pk_pos + 3])
                        if (next_char2 == 3 and next_char3 == 4) or (next_char2 == 5 and next_char3 == 6):
                            logger.error(f"CRITICAL: ZIP signature found at position {pk_pos} in content_text! Rejecting.")
                            content_text = "[Binary file content detected. Text extraction failed.]"
                            content_metadata = {**content_metadata, "extraction_error": "ZIP signature detected"}
            
            # Final UTF-8 encoding validation
            try:
                # Try to encode as UTF-8 to validate (PostgreSQL requires valid UTF-8)
                content_text.encode('utf-8')
            except UnicodeEncodeError as e:
                logger.error(f"CRITICAL: UTF-8 encoding validation failed at final check: {e}. Replacing with error message.")
                content_text = "[Invalid UTF-8 content detected. Text extraction may have failed.]"
                content_metadata = {**content_metadata, "encoding_error": str(e)}
        
        elif isinstance(content_text, bytes):
            # This should NEVER happen, but if it does, reject it
            logger.error(f"CRITICAL: content_text is still bytes at final validation! File: {file_path}")
            content_text = "[Binary content detected. Text extraction failed.]"
            content_metadata = {**content_metadata, "extraction_error": "Binary content at final validation"}
        
        # Ensure content_text is a string (never bytes or None)
        if not isinstance(content_text, str):
            content_text = "[Text extraction failed. Invalid content type.]"
        
        # Final safety: remove any remaining null bytes one last time
        if isinstance(content_text, str) and '\x00' in content_text:
            logger.error("CRITICAL: Null bytes found at VERY final check! This should not happen.")
            content_text = content_text.replace('\x00', '')
            # If ZIP signature appears after null removal, reject
            if len(content_text) >= 4 and content_text.startswith('PK'):
                char2 = ord(content_text[2]) if len(content_text) > 2 else 0
                char3 = ord(content_text[3]) if len(content_text) > 3 else 0
                if (char2 == 3 and char3 == 4) or (char2 == 5 and char3 == 6):
                    content_text = "[Binary content could not be sanitized. Text extraction failed.]"
        
        return {
            "content_text": content_text,
            "content_metadata": content_metadata
        }
    
    except Exception as e:
        logger.error(f"Error extracting text from file: {e}")
        raise


async def process_rag_document(file_path: Optional[str] = None, url: Optional[str] = None, 
                               file_type: Optional[str] = None) -> Dict[str, Any]:
    """Process a RAG document (file or URL) and extract content"""
    try:
        if url:
            # Process URL
            result = await fetch_url_content(url)
            return result
        elif file_path:
            # Process file
            result = await extract_text_from_file(file_path, file_type or "application/octet-stream")
            return result
        else:
            raise ValueError("Either file_path or url must be provided")
    except Exception as e:
        logger.error(f"Error processing RAG document: {e}")
        raise


async def process_and_index_document(
    text_content: str,
    document_id: str,
    law_name: str,
    year: Optional[int] = None,
    authority: Optional[str] = None,
    jurisdiction: str = "Nigeria",
    csv_metadata: Optional[Dict[str, Any]] = None  # NEW: CSV metadata parameter
) -> Dict[str, Any]:
    """
    Complete RAG document processing pipeline:
    1. Convert text to structured Markdown with YAML front-matter
    2. Chunk by legal sections
    3. Generate embeddings
    4. Store in vector database
    
    This is the production-grade processing that enforces structured format.
    """
    try:
        from services.document_processor import process_document_for_rag
        from services.embedding_service import get_embedding_service
        from services.vector_store import get_vector_store
        
        # Get embedding dimension
        embedding_service = get_embedding_service()
        test_embedding = embedding_service.embed_text("test")
        vector_size = len(test_embedding)
        
        # Step 1: Convert to Markdown and chunk
        markdown_content, chunks = process_document_for_rag(
            text_content=text_content,
            law_name=law_name,
            year=year,
            authority=authority,
            jurisdiction=jurisdiction,
            csv_metadata=csv_metadata  # NEW: Pass CSV metadata
        )
        
        if not chunks:
            raise ValueError("No chunks generated from document")
        
        # Step 2: Generate embeddings for all chunks
        embedding_service = get_embedding_service()
        chunk_texts = [chunk['text'] for chunk in chunks]
        embeddings = embedding_service.embed_batch(chunk_texts)
        
        # Step 3: Store in vector database
        vector_store = get_vector_store(vector_size=vector_size)
        vector_store.add_chunks(
            chunks=chunks,
            embeddings=embeddings,
            document_id=document_id
        )
        
        return {
            "markdown_content": markdown_content,
            "chunks_count": len(chunks),
            "processed": True
        }
        
    except Exception as e:
        logger.error(f"Error processing and indexing document: {e}", exc_info=True)
        raise


async def delete_rag_file(file_path: str) -> None:
    """Delete a RAG document file from disk"""
    try:
        if os.path.exists(file_path):
            os.remove(file_path)
            logger.info(f"Deleted file: {file_path}")
    except Exception as e:
        logger.error(f"Error deleting file {file_path}: {e}")
        # Don't raise - file deletion failure shouldn't break the flow
